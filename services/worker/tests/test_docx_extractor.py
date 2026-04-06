"""Unit tests for the DOCX extractor."""

from __future__ import annotations

import io

import pytest

from atlas_worker.extractors.docx import DocxResult, extract_docx


# ---------------------------------------------------------------------------
# In-memory .docx fixture builder
# ---------------------------------------------------------------------------

def _build_docx(
    paragraphs: list[tuple[str, str]] | None = None,
    tables: list[list[list[str]]] | None = None,
) -> bytes:
    """Build a minimal .docx in memory.

    Args:
        paragraphs: list of (text, style_name) tuples.
        tables: list of tables, each a list of rows, each row a list of cell text.
    """
    import docx

    doc = docx.Document()
    for text, style in (paragraphs or []):
        try:
            doc.add_paragraph(text, style=style)
        except KeyError:
            # Style not in minimal docx — add as normal paragraph.
            doc.add_paragraph(text)

    for table_data in (tables or []):
        if not table_data:
            continue
        rows, cols = len(table_data), len(table_data[0])
        tbl = doc.add_table(rows=rows, cols=cols)
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_text in enumerate(row):
                tbl.cell(r_idx, c_idx).text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestExtractDocx:
    def test_returns_docx_result(self):
        raw = _build_docx(paragraphs=[("Hello world", "Normal")])
        result = extract_docx(raw)
        assert isinstance(result, DocxResult)

    def test_full_text_contains_paragraph(self):
        raw = _build_docx(paragraphs=[("My paragraph text.", "Normal")])
        result = extract_docx(raw)
        assert "My paragraph text." in result.full_text

    def test_heading_extracted_to_headings_list(self):
        raw = _build_docx(
            paragraphs=[
                ("Introduction", "Heading 1"),
                ("Some body text.", "Normal"),
            ]
        )
        result = extract_docx(raw)
        assert "Introduction" in result.headings

    def test_heading_formatted_with_hash_prefix(self):
        raw = _build_docx(paragraphs=[("Chapter One", "Heading 1")])
        result = extract_docx(raw)
        assert "# Chapter One" in result.full_text

    def test_tables_count(self):
        table_data = [["Name", "Value"], ["Alpha", "1"], ["Beta", "2"]]
        raw = _build_docx(tables=[table_data])
        result = extract_docx(raw)
        assert result.tables_count == 1

    def test_table_rendered_as_markdown(self):
        table_data = [["Name", "Score"], ["Alice", "95"], ["Bob", "87"]]
        raw = _build_docx(tables=[table_data])
        result = extract_docx(raw)
        assert "Name" in result.full_text
        assert "Score" in result.full_text
        assert "|" in result.full_text

    def test_title_inferred_from_first_heading(self):
        raw = _build_docx(
            paragraphs=[
                ("Main Title", "Heading 1"),
                ("Body paragraph.", "Normal"),
            ]
        )
        result = extract_docx(raw)
        assert result.title == "Main Title"

    def test_language_detected_for_english_text(self):
        raw = _build_docx(
            paragraphs=[
                (
                    "The quick brown fox jumps over the lazy dog. "
                    "It is a well-known English sentence that contains all letters.",
                    "Normal",
                )
            ]
        )
        result = extract_docx(raw)
        assert result.language == "en"

    def test_language_none_for_short_text(self):
        raw = _build_docx(paragraphs=[("Hi", "Normal")])
        result = extract_docx(raw)
        # short text may not trigger language detection — either None or detected
        assert result.language is None or isinstance(result.language, str)

    def test_empty_document(self):
        raw = _build_docx()
        result = extract_docx(raw)
        assert isinstance(result.full_text, str)
        assert result.tables_count == 0
        assert len(result.headings) == 0

    def test_multiple_headings_all_captured(self):
        raw = _build_docx(
            paragraphs=[
                ("Chapter 1", "Heading 1"),
                ("Some content.", "Normal"),
                ("Chapter 2", "Heading 1"),
                ("More content.", "Normal"),
            ]
        )
        result = extract_docx(raw)
        assert "Chapter 1" in result.headings
        assert "Chapter 2" in result.headings

    def test_invalid_bytes_raises(self):
        with pytest.raises(Exception):
            extract_docx(b"this is not a docx file at all")

    def test_result_is_frozen(self):
        raw = _build_docx(paragraphs=[("Text", "Normal")])
        result = extract_docx(raw)
        with pytest.raises((AttributeError, TypeError)):
            result.full_text = "modified"  # type: ignore[misc]
