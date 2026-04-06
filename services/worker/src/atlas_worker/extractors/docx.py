"""DOCX extractor — uses python-docx to extract text and document structure."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field


@dataclass(frozen=True)
class DocxResult:
    full_text: str
    title: str | None
    headings: tuple[str, ...]
    tables_count: int
    language: str | None


def extract_docx(raw: bytes) -> DocxResult:
    """Extract text, headings, and table count from a .docx byte payload.

    Inputs:  raw bytes of a .docx file.
    Outputs: DocxResult with full Markdown-like text, inferred title, heading list,
             table count, and detected language tag.
    Failures: raises ValueError if python-docx cannot open the bytes;
              ImportError if python-docx is not installed.
    """
    import docx  # deferred so the module is importable without the dep installed

    doc = docx.Document(io.BytesIO(raw))

    heading_styles = {"heading 1", "heading 2", "heading 3", "heading 4", "heading 5", "heading 6"}
    lines: list[str] = []
    headings: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower() if para.style else ""
        if style_name in heading_styles:
            level = _heading_level(style_name)
            prefix = "#" * level
            lines.append(f"{prefix} {text}")
            headings.append(text)
        else:
            lines.append(text)

    for table in doc.tables:
        lines.append(_render_table(table))

    full_text = "\n\n".join(lines)
    title = headings[0] if headings else _infer_title(full_text)
    language = _detect_language_hint(full_text)

    return DocxResult(
        full_text=full_text,
        title=title,
        headings=tuple(headings),
        tables_count=len(doc.tables),
        language=language,
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _heading_level(style_name: str) -> int:
    """Return numeric heading level from style name like 'heading 2'."""
    parts = style_name.split()
    try:
        return int(parts[-1])
    except (ValueError, IndexError):
        return 1


def _render_table(table: object) -> str:
    """Render a docx table as a Markdown table string."""
    rows = []
    for row in table.rows:  # type: ignore[union-attr]
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    # Insert markdown separator after header row.
    separator = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"  # type: ignore[union-attr]
    return "\n".join([rows[0], separator] + rows[1:])


# Minimal language detection shared with text.py approach.
_LANG_HINTS: dict[str, frozenset[str]] = {
    "en": frozenset({"the", "and", "of", "to", "in", "is", "it", "that"}),
    "es": frozenset({"de", "la", "el", "en", "que", "y", "los", "las", "un"}),
    "fr": frozenset({"de", "la", "le", "les", "et", "en", "un", "une", "du"}),
    "de": frozenset({"der", "die", "das", "und", "in", "zu", "den", "ist"}),
}


def _detect_language_hint(text: str) -> str | None:
    """Return ISO-639-1 code for the most likely language, or None if uncertain."""
    words = re.findall(r"\b[a-zA-Z]{2,}\b", text.lower())
    if not words:
        return None
    sample = set(words[:500])
    scores = {lang: len(sample & hints) for lang, hints in _LANG_HINTS.items()}
    best_lang, best_score = max(scores.items(), key=lambda x: x[1])
    return best_lang if best_score >= 3 else None


def _infer_title(text: str) -> str | None:
    """Return first non-empty line up to 200 chars as a title candidate."""
    for line in text.splitlines():
        stripped = line.lstrip("#").strip()
        if stripped:
            return stripped[:200] if len(stripped) <= 200 else None
    return None
